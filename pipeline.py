import os
import zipfile
import base64
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.oxml.ns import qn
from openai import OpenAI
from docx.text.paragraph import Paragraph
from docx.table import Table
import re

# ==========================================
# 1. Asset Management (File I/O)
# ==========================================
class AssetExtractor:
    """Handles extracting binary assets (images, embedded files) from the docx archive."""
    def __init__(self, docx_path: str, output_dir: Path):
        self.docx_path = docx_path
        self.media_dir = output_dir / "media"
        self.embed_dir = output_dir / "embeddings"
        
        self.media_dir.mkdir(parents=True, exist_ok=True)
        self.embed_dir.mkdir(parents=True, exist_ok=True)

    def extract_all(self):
        with zipfile.ZipFile(self.docx_path, 'r') as z:
            for file_info in z.infolist():
                if file_info.filename.startswith('word/media/'):
                    filename = os.path.basename(file_info.filename)
                    with open(self.media_dir / filename, "wb") as f:
                        f.write(z.read(file_info.filename))
                elif file_info.filename.startswith('word/embeddings/'):
                    filename = os.path.basename(file_info.filename)
                    with open(self.embed_dir / filename, "wb") as f:
                        f.write(z.read(file_info.filename))


# ==========================================
# 2.1 Text Processing Module
# ==========================================
class TextProcessor:
    """Handles parsing Word paragraphs, extracting styles, runs, and colors into Markdown."""
    
    @staticmethod
    def format_paragraph(paragraph) -> str:
        if not paragraph.text.strip():
            return ""

        # 1. Process Inline Formatting (Bold, Italic, Underline, Color)
        formatted_text = ""
        for run in paragraph.runs:
            raw_text = run.text
            if not raw_text:
                continue
            
            # Separate the leading spaces, the core text, and the trailing spaces
            match = re.match(r'^(\s*)(.*?)(\s*)$', raw_text, re.DOTALL)
            leading_space, core_text, trailing_space = match.groups()

            # If the run was literally just spaces, add them and move on
            if not core_text:
                formatted_text += raw_text
                continue
            
            # Apply native Markdown formatting ONLY to the core text
            if run.bold:
                core_text = f"**{core_text}**"
            if run.italic:
                core_text = f"*{core_text}*"
                
            # Apply HTML fallbacks for Underline and Color
            if run.underline:
                core_text = f"<u>{core_text}</u>"
            
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb)
                core_text = f'<span style="color:#{color_hex}">{core_text}</span>'
                
            # Reassemble the string with the spaces on the OUTSIDE of the formatting
            formatted_text += f"{leading_space}{core_text}{trailing_space}"

        # 2. Process Block-Level Styles (Headings, Titles, Lists)
        style_name = paragraph.style.name.lower()
        
        # Map Word styles to Markdown syntax
        if 'title' in style_name:
            return f"\n# {formatted_text}\n"
        elif 'subtitle' in style_name:
            return f"\n## {formatted_text}\n"
        elif style_name.startswith('heading'):
            # Extracts the number from "Heading 1", "Heading 2", etc.
            level = ''.join(filter(str.isdigit, style_name))
            if level:
                return f"\n{'#' * int(level)} {formatted_text}\n"
            return f"\n# {formatted_text}\n" # Fallback
        elif 'list bullet' in style_name:
            return f"* {formatted_text}\n"
        elif 'list number' in style_name:
            return f"1. {formatted_text}\n" 
        else:
            # Normal text
            return f"\n{formatted_text}\n"

# ==========================================
# 2.2 Table Processing Module
# ==========================================
class TableProcessor:
    """Handles parsing native Word tables and converting them to Markdown tables."""
    
    @staticmethod
    def process_table_to_markdown(table_element, parent_doc) -> str:
        # Wrap the raw XML element in a python-docx Table object
        table = Table(table_element, parent_doc)
        
        if not table.rows:
            return ""

        markdown_table = []
        
        # Helper function to clean cell text
        def clean_cell(cell):
            # Join multiple paragraphs inside a single cell with a space
            text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            # Escape pipe characters so they don't break the Markdown columns
            text = text.replace("|", "\\|")
            # Remove any stray newlines
            text = text.replace("\n", " ").replace("\r", "")
            return text

        # 1. Process Header (Assumes the first row is the header)
        header_cells = [clean_cell(cell) for cell in table.rows[0].cells]
        markdown_table.append("| " + " | ".join(header_cells) + " |")
        
        # 2. Process Separator Row
        # Generates the |---|---| formatting required by Markdown
        separator = ["---"] * len(table.rows[0].cells)
        markdown_table.append("| " + " | ".join(separator) + " |")
        
        # 3. Process Data Rows
        for row in table.rows[1:]:
            row_cells = [clean_cell(cell) for cell in row.cells]
            # Only append the row if it actually contains text
            if any(row_cells):
                markdown_table.append("| " + " | ".join(row_cells) + " |")
            
        return "\n" + "\n".join(markdown_table) + "\n"

# ==========================================
# 3. Vision Processing Module (LLM integration)
# ==========================================
class VisionProcessor:
    """Handles converting images to structured text/code using a Multimodal LLM."""
    def __init__(self, api_client: OpenAI, model_name: str = "gpt-4o"):
        self.client = api_client
        self.model_name = model_name

    def _encode_image(self, image_path: Path) -> str:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def process(self, image_path: Path) -> str:
        if not image_path.exists():
            return f"\n[Error: Image file not found at {image_path}]\n"

        base64_image = self._encode_image(image_path)
        prompt = (
            "Analyze this image from a document.\n"
            "1. If diagram/flowchart: Output valid Mermaid.js code inside a ```mermaid block.\n"
            "2. If table: Output a Markdown table.\n"
            "3. If text: Extract clearly.\n"
            "4. If decorative logo/icon: Reply exactly: \n"
            "Output ONLY the resulting Markdown/Code."
        )

        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }],
                max_tokens=2000
            )
            return f"\n\n{response.choices[0].message.content.strip()}\n\n"
        except Exception as e:
            return f"\n[Vision API Error: {str(e)}]\n"


# ==========================================
# 4. Core Pipeline Orchestrator
# ==========================================
class DocxConversionPipeline:
    """Orchestrates the conversion process based on provided configuration."""
    def __init__(self, docx_path: str, output_dir: str = "output"):
        self.docx_path = docx_path
        self.output_dir = Path(output_dir)
        self.doc = Document(docx_path)
        self.markdown_chunks: List[str] = []
        
        # Initialize sub-modules
        self.assets = AssetExtractor(docx_path, self.output_dir)
        self.vision_processor: Optional[VisionProcessor] = None

    def enable_vision(self, api_client: OpenAI):
        """Injects the Vision Processor into the pipeline."""
        self.vision_processor = VisionProcessor(api_client)

    def run(self):
        """Executes the conversion pipeline."""
        # Step 1: Unpack assets
        self.assets.extract_all()
        
        # Step 2: Traverse document tree
        for element in self.doc.element.body:
            self._route_element(element)
            
        # Step 3: Write final artifact
        self._save_output()

    def _route_element(self, element):
        """Routes XML elements to the appropriate processing module."""
        # 1. Text Paragraphs & Inline Images
        if element.tag.endswith('p'):
            # Correctly wrap the raw XML element into a python-docx Paragraph object
            p = Paragraph(element, self.doc)
            
            # Extract text if present
            if p.text.strip():
                self.markdown_chunks.append(TextProcessor.format_paragraph(p))
                
            # Extract inline images (Word often puts images inside paragraph tags)
            inline_images = element.xpath('.//w:drawing//pic:pic//hlLink | .//w:drawing//pic:pic//a:blip/@r:embed')
            for rId in inline_images:
                self._handle_image(rId)

        # 2. Native Tables
        elif element.tag.endswith('tbl'):
            md_table = TableProcessor.process_table_to_markdown(element, self.doc)
            self.markdown_chunks.append(md_table)

        # 3. Embedded OLE Objects (Excel, Outlook)
        ole_objects = element.xpath('.//*[local-name()="OLEObject"]')
        for obj in ole_objects:
            rId = obj.get(qn('r:id'))
            if rId:
                part = self.doc.part.related_parts[rId]
                filename = os.path.basename(part.partname)
                self.markdown_chunks.append(f"\n[Embedded File: {filename}](./embeddings/{filename})\n")

    def _handle_image(self, rId: str):
        """Delegates image processing based on pipeline configuration."""
        part = self.doc.part.related_parts[rId]
        image_filename = os.path.basename(part.partname)
        image_path = self.assets.media_dir / image_filename

        # If vision is enabled, process it. Otherwise, drop a placeholder.
        if self.vision_processor:
            print(f"Processing image: {image_filename}...")
            result = self.vision_processor.process(image_path)
            self.markdown_chunks.append(result)
        else:
            self.markdown_chunks.append(f"\n![Image Placeholder](./media/{image_filename})\n")

    def _save_output(self, filename="document.md"):
        output_file = self.output_dir / filename
        with open(output_file, "w", encoding="utf-8") as f:
            f.write("".join(self.markdown_chunks))
        print(f"Pipeline complete. File saved to {output_file}")


# ==========================================
# Execution Entry Points
# ==========================================
if __name__ == "__main__":
    DOCX_FILE = "sample_large_doc.docx"
    
    # ---------------------------------------------------------
    # Scenario A: Fast Run (Text and placeholders only)
    # ---------------------------------------------------------
    print("Running Text-Only Pipeline...")
    text_pipeline = DocxConversionPipeline(DOCX_FILE, output_dir="output_fast")
    text_pipeline.run()  # Skips LLM completely
    
    # For my own debugging
    # tmp_doc = Document(DOCX_FILE)
    # for element in tmp_doc.element.body:
    #     if element.tag.endswith('p'):
    #         p = Paragraph(element, tmp_doc)
    #         tmp_txt = TextProcessor.format_paragraph(p)
    #         print(tmp_txt)
    
    # ---------------------------------------------------------
    # Scenario B: Full Pipeline (Text + Vision)
    # ---------------------------------------------------------
    # print("\nRunning Full Pipeline with Vision...")
    # client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    
    # full_pipeline = DocxConversionPipeline(DOCX_FILE, output_dir="output_full")
    # full_pipeline.enable_vision(client) # Dependency injection
    # full_pipeline.run()