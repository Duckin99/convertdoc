import os
from docx import Document
from docx.shared import Inches, RGBColor
from PIL import Image, ImageDraw

def create_dummy_image(filename="test_diagram.png"):
    """Creates a simple test image to embed in the document."""
    img = Image.new('RGB', (400, 200), color=(73, 109, 137))
    d = ImageDraw.Draw(img)
    d.text((50, 90), "MOCK DIAGRAM IMAGE", fill=(255, 255, 0))
    img.save(filename)
    return filename

def generate_messy_mock_docx(output_filename="sample_large_doc.docx"):
    doc = Document()

    # 1. Title and Subtitle
    doc.add_heading('Messy Technical Specification', 0)
    doc.add_paragraph('Subtitle: Edge Case Testing', style='Subtitle')

    # 2. Messy Inline Text Formatting
    doc.add_heading('1. Text Formatting Nightmare', level=1)
    
    p = doc.add_paragraph('This paragraph tests standard text, followed by ')
    p.add_run('bold text').bold = True
    p.add_run(', then ')
    
    # Testing the space bug: italic text with trailing spaces
    italic_run = p.add_run('italic text with spaces   ')
    italic_run.italic = True
    
    p.add_run('and finally ')
    
    # Testing Underline
    underline_run = p.add_run('underlined text')
    underline_run.underline = True
    p.add_run('.')

    # Testing Colors
    p2 = doc.add_paragraph('Here is a warning message in ')
    red_run = p2.add_run('bright red text')
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    p2.add_run(' to test HTML span tags.')

    # 3. Lists
    doc.add_heading('2. Requirements Lists', level=1)
    doc.add_paragraph('Backend Requirements:', style='Heading 3')
    doc.add_paragraph('Setup PostgreSQL database', style='List Bullet')
    doc.add_paragraph('Configure Redis cache', style='List Bullet')
    
    doc.add_paragraph('Deployment Steps:', style='Heading 3')
    doc.add_paragraph('Build Docker image', style='List Number')
    doc.add_paragraph('Deploy to Kubernetes', style='List Number')

    # 4. Add an Image
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph('Below is a placeholder image representing a Mermaid diagram:')
    img_path = create_dummy_image()
    doc.add_picture(img_path, width=Inches(3.0))
    os.remove(img_path)

    # 5. A Truly Messy Table
    doc.add_heading('4. Messy API Endpoints', level=1)
    doc.add_paragraph('This table contains line breaks and pipe characters inside the cells, which usually destroys Markdown tables:')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Endpoint'
    hdr_cells[2].text = 'Description / Notes'
    
    # Row 1 (Standard)
    row_cells = table.rows[1].cells
    row_cells[0].text = 'GET'
    row_cells[1].text = '/api/v1/users'
    row_cells[2].text = 'Fetches all users. Standard behavior.'
    
    # Row 2 (The Nightmare Cell)
    row_cells = table.rows[2].cells
    row_cells[0].text = 'POST'
    row_cells[1].text = '/api/v1/users'
    
    # Simulating a user pressing 'Enter' multiple times inside a single Word table cell
    # And adding a '|' character to see if our escape logic works
    nightmare_cell = row_cells[2]
    nightmare_cell.text = "Creates a user.\n\nWARNING: Must include a | character in the payload.\n\nSee docs for details."

    # Save the final file
    doc.save(output_filename)
    print(f"Successfully created '{output_filename}' with messy edge cases!")

if __name__ == "__main__":
    generate_messy_mock_docx()